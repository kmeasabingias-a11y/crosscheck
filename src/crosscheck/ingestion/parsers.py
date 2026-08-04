"""Document parsing — the pipeline's front door.

``parse(path)`` dispatches on file extension to a format-specific parser, each of which
returns a :class:`~crosscheck.models.Document` (spec v2 §7.1): an ordered list of
:class:`~crosscheck.models.Section` objects carrying headings and, where the format
exposes them, page spans. Four formats ship in v1: PDF (text-based), DOCX, Markdown,
and plain text. Scanned/image PDFs are out of scope (§3) — only extractable text is read.

Design notes:
- The ``doc_id`` is a content hash of the document text, so re-ingesting the same file
is idempotent (§4 resume) and byte-identical duplicates collapse to one document.
- PDFs have no reliable heading markup, so each page becomes one section and a
repetition heuristic strips running headers/footers and page-number lines (§7.1).
- Markdown and DOCX carry real structure (ATX/setext headings, heading styles); plain
text has none, so it is a single section and the chunker does all the splitting.
"""

import re
from collections import Counter
from collections.abc import Callable
from pathlib import Path
from typing import Any, NamedTuple

import pdfplumber
from docx import Document as DocxDocument
from loguru import logger
from markdown_it import MarkdownIt

from crosscheck.ids import doc_id, section_id
from crosscheck.models import Document, Section


class UnsupportedFormatError(ValueError):
    """Raised when :func:`parse` is given a file whose extension it cannot handle."""


class _RawSection(NamedTuple):
    """A section before ids are assigned: heading, body text, and optional page span."""

    heading: str | None
    text: str
    page_span: tuple[int, int] | None = None


# A standalone line that is only a page number, e.g. "12", "Page 12", or "12 / 340".
_PAGE_NUMBER_RE = re.compile(r"^\s*(?:page\s+)?\d+(?:\s*/\s*\d+)?\s*$", re.IGNORECASE)
# DOCX paragraph styles that start a new section: "Heading 1".."Heading 9", "Title", "Subtitle".
_HEADING_STYLE_RE = re.compile(r"^(?:heading\s*\d+|title|subtitle)$", re.IGNORECASE)


def _assemble(
    path: Path,
    raw: list[_RawSection],
    *,
    title: str | None,
    metadata: dict[str, Any],
) -> Document:
    """Assign ids and build the final :class:`Document` from raw sections.

    Empty sections (no non-whitespace text) are dropped. The ``doc_id`` is a content
    hash of the surviving section text, so re-ingesting the same file is idempotent and
    byte-identical duplicates collapse to one document; section ids are position-based.
    """
    kept = [section for section in raw if section.text.strip()]
    full_text = "\n\n".join(section.text for section in kept)
    did = doc_id(full_text)
    sections = [
        Section(
            section_id=section_id(did, ordinal),
            heading=section.heading,
            text=section.text,
            page_span=section.page_span,
        )
        for ordinal, section in enumerate(kept)
    ]
    if not sections:
        logger.warning("parser: {} produced no non-empty sections", path.name)
    return Document(
        doc_id=did,
        source_path=path,
        title=title,
        sections=sections,
        metadata=metadata,
    )


def _parse_txt(path: Path) -> Document:
    """Parse a plain-text file as a single section (the chunker does all splitting)."""
    text = path.read_text(encoding="utf-8", errors="replace")
    return _assemble(
        path,
        [_RawSection(heading=None, text=text)],
        title=path.stem,
        metadata={"format": "txt"},
    )


def _parse_markdown(path: Path) -> Document:
    """Parse Markdown, splitting into sections on ATX/setext headings.

    Uses markdown-it's block-token line maps to slice the original source, so each
    section's body is the verbatim text between one heading and the next (the heading
    line itself is captured in :attr:`Section.heading`, not repeated in the body).
    """
    source = path.read_text(encoding="utf-8", errors="replace")
    lines = source.splitlines()
    tokens = MarkdownIt().parse(source)

    # (start_line, end_line, heading_text) for every heading, in document order.
    headings: list[tuple[int, int, str]] = []
    for index, token in enumerate(tokens):
        if token.type == "heading_open" and token.map is not None:
            inline = tokens[index + 1] if index + 1 < len(tokens) else None
            heading_text = (
                inline.content.strip() if inline is not None and inline.type == "inline" else ""
            )
            headings.append((token.map[0], token.map[1], heading_text))

    raw: list[_RawSection] = []
    if not headings:
        raw.append(_RawSection(heading=None, text=source))
    else:
        if headings[0][0] > 0:  # preamble before the first heading
            raw.append(_RawSection(heading=None, text="\n".join(lines[: headings[0][0]])))
        for index, (_start, end, heading_text) in enumerate(headings):
            body_end = headings[index + 1][0] if index + 1 < len(headings) else len(lines)
            raw.append(
                _RawSection(heading=heading_text or None, text="\n".join(lines[end:body_end]))
            )

    title = headings[0][2] if headings and headings[0][2] else path.stem
    return _assemble(path, raw, title=title, metadata={"format": "markdown"})


def _parse_docx(path: Path) -> Document:
    """Parse a DOCX file, splitting into sections on heading-styled paragraphs."""
    document = DocxDocument(str(path))
    raw: list[_RawSection] = []
    heading: str | None = None
    body: list[str] = []
    doc_title: str | None = None

    for paragraph in document.paragraphs:
        style = (paragraph.style.name if paragraph.style is not None else "").strip()
        text = paragraph.text.strip()
        if _HEADING_STYLE_RE.match(style):
            if body:
                raw.append(_RawSection(heading, "\n".join(body)))
                body = []
            heading = text or None
            if doc_title is None and style.lower() in {"title", "heading 1"}:
                doc_title = heading
        elif text:
            body.append(text)
    if body:
        raw.append(_RawSection(heading, "\n".join(body)))

    core_title = document.core_properties.title or None
    title = doc_title or core_title or path.stem
    return _assemble(path, raw, title=title, metadata={"format": "docx"})


def _strip_running_headers_footers(
    pages: list[str],
    *,
    edge_lines: int = 2,
    min_fraction: float = 0.5,
) -> list[str]:
    """Remove running headers/footers and page-number lines from page texts.

    A non-empty line that appears among the top or bottom ``edge_lines`` of at least
    ``min_fraction`` of the pages is treated as a running header/footer and removed from
    every page; standalone page-number lines are always removed (spec v2 §7.1). With
    fewer than three pages there is too little signal to infer repetition, so only page
    numbers are stripped.

    Args:
        pages: The extracted text of each page, in order.
        edge_lines: How many lines at the top and bottom of a page to inspect.
        min_fraction: Fraction of pages a line must appear on to count as running.

    Returns:
        The page texts with running headers/footers and page numbers removed.
    """
    running: set[str] = set()
    if len(pages) >= 3:
        edge_counts: Counter[str] = Counter()
        for text in pages:
            page_lines = [line.strip() for line in text.splitlines() if line.strip()]
            edges = page_lines[:edge_lines] + page_lines[-edge_lines:]
            for line in set(edges):
                edge_counts[line] += 1
        threshold = max(2, int(len(pages) * min_fraction))
        running = {line for line, count in edge_counts.items() if count >= threshold}

    cleaned: list[str] = []
    for text in pages:
        kept: list[str] = []
        for line in text.splitlines():
            stripped = line.strip()
            if stripped and (stripped in running or _PAGE_NUMBER_RE.match(stripped)):
                continue
            kept.append(line)
        cleaned.append("\n".join(kept))
    return cleaned


def _parse_pdf(path: Path) -> Document:
    """Parse a text-based PDF: one section per page, headers/footers stripped.

    PDFs carry no reliable heading markup, so each page becomes a :class:`Section`
    with a 1-based ``page_span``. Scanned/image PDFs are out of scope (spec §3).
    """
    pages: list[str] = []
    pdf_title: Any = None
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
        pdf_title = (pdf.metadata or {}).get("Title")

    cleaned = _strip_running_headers_footers(pages)
    raw = [
        _RawSection(heading=None, text=text, page_span=(number, number))
        for number, text in enumerate(cleaned, start=1)
    ]
    title = (
        pdf_title.strip() if isinstance(pdf_title, str) and pdf_title.strip() else None
    ) or path.stem
    return _assemble(path, raw, title=title, metadata={"format": "pdf", "page_count": len(pages)})


_PARSERS: dict[str, Callable[[Path], Document]] = {
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt": _parse_txt,
    ".text": _parse_txt,
}

#: Extensions :func:`parse` can handle (§3 ships PDF, DOCX, Markdown and plain text). Public
#: so a caller can filter *before* dispatching: the API skips unsupported uploads rather than
#: failing an otherwise good multi-file request, and it should not have to catch an exception
#: per file to find that out.
SUPPORTED_SUFFIXES: frozenset[str] = frozenset(_PARSERS)


def parse(path: Path) -> Document:
    """Parse a document file into a :class:`Document`, dispatching on its extension.

    Supported formats: PDF (text-based), DOCX, Markdown (``.md`` / ``.markdown``), and
    plain text (``.txt`` / ``.text``). Scanned/image PDFs are out of scope (spec §3).

    Args:
        path: Path to the file to parse.

    Returns:
        A :class:`Document` with an ordered list of :class:`Section` objects.

    Raises:
        FileNotFoundError: If ``path`` does not exist.
        UnsupportedFormatError: If the file extension is not supported.
    """
    if not path.exists():
        raise FileNotFoundError(path)
    parser = _PARSERS.get(path.suffix.lower())
    if parser is None:
        supported = ", ".join(sorted(_PARSERS))
        raise UnsupportedFormatError(
            f"unsupported file type {path.suffix!r} for {path}; supported: {supported}"
        )
    logger.info("parsing {} ({})", path.name, path.suffix.lower())
    document = parser(path)
    logger.info(
        "parsed {} into {} section(s), doc_id={}",
        path.name,
        len(document.sections),
        document.doc_id,
    )
    return document

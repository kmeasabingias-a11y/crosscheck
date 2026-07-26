"""Unit tests for the document parsers."""

from pathlib import Path

import pytest
from docx import Document as DocxDocument
from fpdf import FPDF

from crosscheck.ingestion.parsers import (
    UnsupportedFormatError,
    _strip_running_headers_footers,
    parse,
)


def test_txt_is_single_section(tmp_path: Path) -> None:
    p = tmp_path / "note.txt"
    p.write_text("Employees receive 20 PTO days.\nContractors do not.", encoding="utf-8")
    doc = parse(p)
    assert doc.metadata["format"] == "txt"
    assert doc.title == "note"
    assert len(doc.sections) == 1
    assert doc.sections[0].heading is None
    assert "20 PTO days" in doc.sections[0].text


def test_doc_id_is_content_addressed(tmp_path: Path) -> None:
    a = tmp_path / "a.txt"
    b = tmp_path / "b_different_name.txt"
    a.write_text("identical body text", encoding="utf-8")
    b.write_text("identical body text", encoding="utf-8")
    assert parse(a).doc_id == parse(b).doc_id  # same content -> same id, regardless of name
    c = tmp_path / "c.txt"
    c.write_text("different body text", encoding="utf-8")
    assert parse(c).doc_id != parse(a).doc_id


def test_markdown_splits_on_headings(tmp_path: Path) -> None:
    p = tmp_path / "policy.md"
    p.write_text(
        "Intro paragraph.\n\n# Leave Policy\nEmployees receive 20 PTO days.\n\n"
        "## Contractors\nContractors are not entitled to PTO.\n",
        encoding="utf-8",
    )
    doc = parse(p)
    assert doc.title == "Leave Policy"
    headings = [s.heading for s in doc.sections]
    assert headings == [None, "Leave Policy", "Contractors"]
    # Heading line is not duplicated inside the section body.
    leave = next(s for s in doc.sections if s.heading == "Leave Policy")
    assert "# Leave Policy" not in leave.text
    assert "20 PTO days" in leave.text


def test_markdown_without_headings_is_one_section(tmp_path: Path) -> None:
    p = tmp_path / "flat.md"
    p.write_text("Just a paragraph with no headings at all.", encoding="utf-8")
    doc = parse(p)
    assert len(doc.sections) == 1
    assert doc.title == "flat"


def test_docx_splits_on_heading_styles(tmp_path: Path) -> None:
    p = tmp_path / "handbook.docx"
    d = DocxDocument()
    d.add_heading("Overview", level=1)
    d.add_paragraph("Employees receive 20 PTO days.")
    d.add_heading("Exceptions", level=2)
    d.add_paragraph("Contractors are not entitled to PTO.")
    d.save(str(p))

    doc = parse(p)
    assert doc.title == "Overview"
    assert [s.heading for s in doc.sections] == ["Overview", "Exceptions"]
    assert "20 PTO days" in doc.sections[0].text
    assert "not entitled" in doc.sections[1].text


def test_unsupported_extension_raises(tmp_path: Path) -> None:
    p = tmp_path / "data.csv"
    p.write_text("a,b,c", encoding="utf-8")
    with pytest.raises(UnsupportedFormatError):
        parse(p)


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        parse(tmp_path / "nope.md")


def test_strip_running_headers_and_page_numbers() -> None:
    pages = [
        "ACME CONFIDENTIAL\nBody of page one.\n1",
        "ACME CONFIDENTIAL\nBody of page two.\n2",
        "ACME CONFIDENTIAL\nBody of page three.\n3",
        "ACME CONFIDENTIAL\nBody of page four.\n4",
    ]
    cleaned = _strip_running_headers_footers(pages)
    joined = "\n".join(cleaned)
    assert "ACME CONFIDENTIAL" not in joined  # repeated header stripped
    assert not any(line.strip().isdigit() for line in joined.splitlines())  # page numbers gone
    assert "Body of page one." in cleaned[0]  # real content preserved


def test_strip_is_conservative_for_few_pages() -> None:
    pages = ["HEADER\nreal content\n1", "HEADER\nother content\n2"]
    cleaned = _strip_running_headers_footers(pages)
    # Too few pages to infer a running header, so HEADER stays; page numbers still go.
    assert cleaned[0].splitlines() == ["HEADER", "real content"]


def _make_pdf(path: Path, pages: list[str], *, title: str | None = None) -> Path:
    """Write a simple multi-page PDF, one page per element of `pages`."""
    pdf = FPDF(format="A4", unit="mm")
    pdf.set_margins(20, 20, 20)
    pdf.set_auto_page_break(auto=True, margin=20)
    if title is not None:
        pdf.set_title(title)
    for text in pages:
        pdf.add_page()
        pdf.set_font("helvetica", "", 11)
        pdf.multi_cell(0, 6, text)
    pdf.output(str(path))
    return path


def test_pdf_makes_one_section_per_page(tmp_path: Path) -> None:
    p = _make_pdf(
        tmp_path / "policy.pdf",
        [
            "Employees receive 20 PTO days per year.",
            "Vendors must carry liability insurance.",
        ],
    )
    doc = parse(p)
    assert doc.metadata["format"] == "pdf"
    assert doc.metadata["page_count"] == 2
    assert len(doc.sections) == 2
    assert doc.sections[0].page_span == (1, 1)
    assert doc.sections[1].page_span == (2, 2)
    assert "20 PTO days" in doc.sections[0].text
    assert "liability insurance" in doc.sections[1].text
    # PDFs carry no reliable heading markup, so sections are unheaded by design.
    assert all(section.heading is None for section in doc.sections)


def test_pdf_title_prefers_metadata_then_falls_back_to_stem(tmp_path: Path) -> None:
    titled = _make_pdf(tmp_path / "a.pdf", ["Some body text."], title="Security Policy v5.2")
    assert parse(titled).title == "Security Policy v5.2"

    untitled = _make_pdf(tmp_path / "handbook_v4.pdf", ["Some body text."])
    assert parse(untitled).title == "handbook_v4"


def test_pdf_strips_running_header_across_pages(tmp_path: Path) -> None:
    header = "ARDEN SYSTEMS INTERNAL"
    p = _make_pdf(
        tmp_path / "standards.pdf",
        [
            f"{header}\nPasswords must be at least 14 characters.",
            f"{header}\nLogs are retained for 13 months.",
            f"{header}\nRemote access requires a compliant device.",
        ],
    )
    doc = parse(p)
    body = "\n".join(section.text for section in doc.sections)
    assert header not in body  # repeated across 3 pages, so stripped
    assert "14 characters" in body
    assert "13 months" in body

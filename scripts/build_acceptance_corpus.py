"""Render the binary documents of the acceptance corpus (spec v2 §8, Phase 3).

The acceptance corpus lives in ``benchmarks/acceptance/corpus/``. Its Markdown and
plain-text documents are committed directly; the DOCX and PDF documents are rendered here
from the JSON sources in ``benchmarks/acceptance/sources/``.

Two things drive that split:

1. **No duplicate documents.** If a source file sat inside the corpus directory beside its
    rendered twin, the audit would parse both and compare a document against itself. The
    sources live in a sibling directory, so the corpus holds exactly one file per document.
2. **Reviewable prose.** A committed ``.docx``/``.pdf`` is opaque to ``git diff``. Keeping
    the text in JSON makes every planted contradiction reviewable, and the binaries
    reproducible from it.

PDF sources tag each section with a ``page``; sections are grouped onto pages in order.
The PDFs are laid out over three pages with a running header and a page-number footer
because ``_strip_running_headers_footers`` only infers a running header from three or more
pages — a shorter fixture would leave that path unexercised on real input.

Run from anywhere::

    uv run python scripts/build_acceptance_corpus.py
"""

import json
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from fpdf import FPDF

_ROOT = Path(__file__).resolve().parent.parent / "benchmarks" / "acceptance"
_SOURCE_DIR = _ROOT / "sources"
_CORPUS_DIR = _ROOT / "corpus"


class _PolicyPDF(FPDF):
    """A policy-document PDF with a running header and a page-number footer.

    The running header is what makes the generated corpus exercise the repetition
    heuristic in ``crosscheck.ingestion.parsers._strip_running_headers_footers``.
    """

    def __init__(self, running_header: str) -> None:
        super().__init__(format="A4", unit="mm")
        self.running_header = running_header
        self.set_margins(25, 20, 25)
        self.set_auto_page_break(auto=True, margin=20)

    def header(self) -> None:
        """Draw the running header at the top of every page."""
        self.set_font("helvetica", "", 8)
        self.cell(0, 6, self.running_header, align="C")
        self.ln(10)

    def footer(self) -> None:
        """Draw the page number at the bottom of every page."""
        self.set_y(-15)
        self.set_font("helvetica", "", 8)
        self.cell(0, 6, str(self.page_no()), align="C")


def _write_docx(spec: dict[str, Any], path: Path) -> None:
    """Render a DOCX document from its JSON source.

    ``_parse_docx`` splits on paragraphs styled "Heading N"/"Title", so headings are
    written with ``add_heading`` rather than as bold body text.

    Args:
        spec: The parsed JSON source.
        path: Destination file.
    """
    document = DocxDocument()
    document.core_properties.title = spec["title"]
    for index, section in enumerate(spec["sections"]):
        document.add_heading(section["heading"], level=1 if index == 0 else 2)
        for paragraph in section["paragraphs"]:
            document.add_paragraph(paragraph)
    document.save(str(path))


def _write_pdf(spec: dict[str, Any], path: Path) -> None:
    """Render a multi-page PDF from its JSON source, breaking pages on the ``page`` tag.

    Args:
        spec: The parsed JSON source.
        path: Destination file.
    """
    pdf = _PolicyPDF(spec["running_header"])
    pdf.set_title(spec["title"])
    current_page: int | None = None
    for section in spec["sections"]:
        if section["page"] != current_page:
            pdf.add_page()
            current_page = section["page"]
        pdf.set_font("helvetica", "B", 12)
        pdf.multi_cell(0, 6, section["heading"])
        pdf.ln(2)
        pdf.set_font("helvetica", "", 10)
        for paragraph in section["paragraphs"]:
            pdf.multi_cell(0, 5, paragraph)
            pdf.ln(2)
        pdf.ln(3)
    pdf.output(str(path))


def main() -> None:
    """Render every JSON source into the corpus directory, overwriting existing copies."""
    _CORPUS_DIR.mkdir(parents=True, exist_ok=True)
    for source in sorted(_SOURCE_DIR.glob("*.json")):
        spec = json.loads(source.read_text(encoding="utf-8"))
        path = _CORPUS_DIR / spec["filename"]
        if spec["format"] == "docx":
            _write_docx(spec, path)
        else:
            _write_pdf(spec, path)
        print(f"wrote {path.relative_to(_ROOT.parent.parent)}")


if __name__ == "__main__":
    main()
